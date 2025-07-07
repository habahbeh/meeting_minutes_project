# ===========================================
# meetings/services.py
# ===========================================

import os
import openai
import json
from langchain.chains.summarize import load_summarize_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain.schema import Document
from django.conf import settings
from django.utils import timezone
from .models import MeetingTask, MeetingDecision


class MeetingProcessor:
    """معالج الاجتماعات"""

    def __init__(self):
        self.openai_client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        self.llm = ChatOpenAI(
            model="gpt-4o",
            openai_api_key=settings.OPENAI_API_KEY,
            temperature=0.3
        )

    def audio_to_text(self, audio_file_path):
        """تحويل الصوت إلى نص باستخدام Whisper"""
        try:
            with open(audio_file_path, "rb") as audio_file:
                transcript = self.openai_client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file,
                    language="ar"  # يمكن تركه فارغ للكشف التلقائي
                )
            return transcript.text
        except Exception as e:
            raise Exception(f"خطأ في تحويل الصوت إلى نص: {str(e)}")

    def create_summary(self, text):
        """إنشاء تلخيص للنص"""
        try:
            # تقسيم النص إلى أجزاء صغيرة
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=2000,
                chunk_overlap=200
            )

            docs = [Document(page_content=text)]
            split_docs = text_splitter.split_documents(docs)

            # قالب التلخيص
            summary_prompt = PromptTemplate(
                template="""
                قم بتلخيص هذا الجزء من محادثة الاجتماع بطريقة واضحة ومفيدة:

                {text}

                التلخيص:
                """,
                input_variables=["text"]
            )

            # إنشاء سلسلة التلخيص
            summary_chain = load_summarize_chain(
                self.llm,
                chain_type="map_reduce",
                map_prompt=summary_prompt,
                combine_prompt=summary_prompt
            )

            summary = summary_chain.run(split_docs)
            return summary

        except Exception as e:
            raise Exception(f"خطأ في إنشاء التلخيص: {str(e)}")

    def create_minutes(self, text):
        """إنشاء محضر اجتماع منظم"""
        try:
            minutes_prompt = f"""
            قم بإنشاء محضر اجتماع منظم ومفصل من النص التالي. يجب أن يكون المحضر بتنسيق رسمي ومهني، ولا يظهر أنه مُنشأ بواسطة AI.

            {text}

            يجب أن يتضمن المحضر الأقسام التالية:

            # محضر اجتماع

            ## معلومات الاجتماع
            - التاريخ: [استخرج التاريخ من المحتوى إن وجد، وإلا استخدم تاريخ اليوم]
            - الحضور: [استخرج قائمة الحاضرين]
            - موضوع الاجتماع: [استخرج موضوع الاجتماع]

            ## جدول الأعمال
            [قائمة بالنقاط الرئيسية التي تمت مناقشتها في الاجتماع]

            ## المناقشات والملاحظات
            [تلخيص منظم للمناقشات التي تمت خلال الاجتماع]

            ## النتائج والتوصيات
            [تلخيص النتائج والتوصيات التي تم التوصل إليها]

            ## الخطوات التالية
            [خطة العمل للمتابعة]

            اكتب المحضر بلغة رسمية وواضحة باللغة العربية. تأكد من أن المحضر يبدو كأنه تم كتابته بواسطة إنسان وليس AI.
            """

            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": minutes_prompt}],
                temperature=0.3,
                max_tokens=2000
            )

            return response.choices[0].message.content

        except Exception as e:
            raise Exception(f"خطأ في إنشاء محضر الاجتماع: {str(e)}")

    def extract_tasks_and_decisions(self, meeting):
        """استخراج المهام والقرارات من محضر الاجتماع"""
        try:
            # جمع النص من المحضر والتلخيص والنص المحول
            all_text = f"""محضر الاجتماع:
            {meeting.minutes}

            تلخيص الاجتماع:
            {meeting.summary}

            النص المحول:
            {meeting.transcript}
            """

            extraction_prompt = f"""
            استخرج المهام والقرارات من محتوى الاجتماع التالي:

            {all_text}

            قم بإخراج النتائج بتنسيق JSON كما يلي:

            {{
                "tasks": [
                    {{
                        "description": "وصف المهمة",
                        "assigned_to": "اسم الشخص المسؤول (إن وجد)",
                        "due_date": "تاريخ الاستحقاق بصيغة YYYY-MM-DD (إن وجد)"
                    }}
                ],
                "decisions": [
                    {{
                        "description": "نص القرار",
                        "decision_number": "رقم القرار (إن وجد)"
                    }}
                ]
            }}

            أخرج JSON فقط بدون أي نص إضافي.
            """

            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": extraction_prompt}],
                temperature=0.2,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )

            result = json.loads(response.choices[0].message.content)

            # حفظ المهام المستخرجة
            for task_data in result.get('tasks', []):
                due_date = None
                if task_data.get('due_date'):
                    try:
                        due_date = timezone.datetime.strptime(task_data['due_date'], '%Y-%m-%d').date()
                    except:
                        pass

                MeetingTask.objects.create(
                    meeting=meeting,
                    description=task_data['description'],
                    assigned_to=task_data.get('assigned_to', ''),
                    due_date=due_date
                )

            # حفظ القرارات المستخرجة
            for decision_data in result.get('decisions', []):
                MeetingDecision.objects.create(
                    meeting=meeting,
                    description=decision_data['description'],
                    decision_number=decision_data.get('decision_number', '')
                )

            return True, "تم استخراج المهام والقرارات بنجاح"

        except Exception as e:
            return False, f"خطأ في استخراج المهام والقرارات: {str(e)}"

    def process_meeting(self, meeting):
        """معالجة الاجتماع كاملة"""
        try:
            # 1. تحويل الصوت إلى نص
            transcript = self.audio_to_text(meeting.audio_file.path)
            meeting.transcript = transcript
            meeting.save()

            # 2. إنشاء التلخيص
            summary = self.create_summary(transcript)
            meeting.summary = summary
            meeting.save()

            # 3. إنشاء المحضر
            minutes = self.create_minutes(transcript)
            meeting.minutes = minutes
            meeting.save()

            # 4. استخراج المهام والقرارات
            self.extract_tasks_and_decisions(meeting)

            return True, "تم معالجة الاجتماع بنجاح"

        except Exception as e:
            return False, f"خطأ في معالجة الاجتماع: {str(e)}"


class MeetingExporter:
    """تصدير محضر الاجتماع بصيغ مختلفة"""

    @staticmethod
    def export_to_pdf(meeting):
        """تصدير المحضر بصيغة PDF"""
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_RIGHT
            from io import BytesIO
            import os

            # تسجيل الخط
            try:
                # تجربة تسجيل خط عربي مناسب
                font_paths = [
                    os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'static', 'fonts', 'Arial.ttf'),
                    '/Library/Fonts/Arial.ttf',  # macOS
                    '/usr/share/fonts/TTF/Arial.ttf',  # Linux
                    'C:\\Windows\\Fonts\\Arial.ttf',  # Windows
                    os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'static', 'fonts',
                                 'Amiri-Regular.ttf'),
                    '/Library/Fonts/Amiri-Regular.ttf',
                    '/usr/share/fonts/TTF/Amiri-Regular.ttf',
                    'C:\\Windows\\Fonts\\Amiri-Regular.ttf',
                ]

                # البحث عن أول خط متوفر
                font_found = False
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        if 'Amiri' in font_path:
                            pdfmetrics.registerFont(TTFont('Arabic', font_path))
                        else:
                            pdfmetrics.registerFont(TTFont('Arabic', font_path))
                        font_found = True
                        break

                if not font_found:
                    # استخدام خط Helvetica كخط بديل إذا لم يتم العثور على أي خط عربي
                    print("لم يتم العثور على خط عربي مناسب، سيتم استخدام خط افتراضي.")
            except Exception as font_error:
                print(f"خطأ في تسجيل الخط: {str(font_error)}")

            # إنشاء أنماط النص
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='Arabic',
                fontName='Arabic',
                alignment=TA_RIGHT,
                fontSize=12,
                leading=14,
                spaceAfter=6
            ))
            styles.add(ParagraphStyle(
                name='ArabicTitle',
                fontName='Arabic',
                alignment=TA_RIGHT,
                fontSize=16,
                leading=18,
                spaceAfter=12,
                bold=True
            ))

            # إنشاء البيانات
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)

            # دالة لمعالجة النص العربي
            def arabic_text(text):
                try:
                    reshaped_text = arabic_reshaper.reshape(text)
                    bidi_text = get_display(reshaped_text)
                    return bidi_text
                except:
                    return text

            # بناء المستند
            content = []

            # العنوان
            title_text = arabic_text(meeting.title)
            content.append(Paragraph(title_text, styles['ArabicTitle']))
            content.append(Spacer(1, 12))

            # التاريخ
            date_text = arabic_text(f"تاريخ الاجتماع: {meeting.created_at.strftime('%Y-%m-%d')}")
            content.append(Paragraph(date_text, styles['Arabic']))
            content.append(Spacer(1, 12))

            # المحضر
            content.append(Paragraph(arabic_text("محضر الاجتماع:"), styles['ArabicTitle']))

            # تقسيم المحضر إلى فقرات
            for line in meeting.minutes.split('\n'):
                if line.strip():
                    line_text = arabic_text(line.strip())
                    content.append(Paragraph(line_text, styles['Arabic']))
                    content.append(Spacer(1, 6))

            # إضافة المهام إذا وجدت
            if hasattr(meeting, 'tasks') and meeting.tasks.exists():
                content.append(Spacer(1, 12))
                content.append(Paragraph(arabic_text("المهام:"), styles['ArabicTitle']))

                for task in meeting.tasks.all():
                    task_text = f"• {task.description}"
                    if task.assigned_to:
                        task_text += f" (مسند إلى: {task.assigned_to})"
                    if task.due_date:
                        task_text += f" (تاريخ الاستحقاق: {task.due_date})"

                    content.append(Paragraph(arabic_text(task_text), styles['Arabic']))

            # إضافة القرارات إذا وجدت
            if hasattr(meeting, 'decisions') and meeting.decisions.exists():
                content.append(Spacer(1, 12))
                content.append(Paragraph(arabic_text("القرارات:"), styles['ArabicTitle']))

                for decision in meeting.decisions.all():
                    decision_text = f"• {decision.description}"
                    if decision.decision_number:
                        decision_text += f" (رقم: {decision.decision_number})"

                    content.append(Paragraph(arabic_text(decision_text), styles['Arabic']))

            # بناء المستند
            doc.build(content)

            # إرجاع البيانات
            pdf_data = buffer.getvalue()
            buffer.close()

            return pdf_data

        except Exception as e:
            raise Exception(f"خطأ في تصدير المحضر كملف PDF: {str(e)}")

    @staticmethod
    def export_to_docx(meeting):
        """تصدير المحضر بصيغة DOCX"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO

            # إنشاء مستند جديد
            doc = Document()

            # إضافة العنوان
            title = doc.add_heading(meeting.title, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # إضافة التاريخ
            date_para = doc.add_paragraph(f"تاريخ الاجتماع: {meeting.created_at.strftime('%Y-%m-%d')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # إضافة المحضر
            minutes_heading = doc.add_heading("محضر الاجتماع:", level=2)
            minutes_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # تقسيم المحضر إلى فقرات
            for line in meeting.minutes.split('\n'):
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # إضافة المهام إذا وجدت
            tasks = meeting.tasks.all()
            if tasks.exists():
                tasks_heading = doc.add_heading("المهام:", level=2)
                tasks_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for task in tasks:
                    task_text = f"• {task.description}"
                    if task.assigned_to:
                        task_text += f" (مسند إلى: {task.assigned_to})"
                    if task.due_date:
                        task_text += f" (تاريخ الاستحقاق: {task.due_date})"

                    task_para = doc.add_paragraph(task_text)
                    task_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # إضافة القرارات إذا وجدت
            decisions = meeting.decisions.all()
            if decisions.exists():
                decisions_heading = doc.add_heading("القرارات:", level=2)
                decisions_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for decision in decisions:
                    decision_text = f"• {decision.description}"
                    if decision.decision_number:
                        decision_text += f" (رقم: {decision.decision_number})"

                    decision_para = doc.add_paragraph(decision_text)
                    decision_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # حفظ المستند
            buffer = BytesIO()
            doc.save(buffer)
            docx_data = buffer.getvalue()
            buffer.close()

            return docx_data

        except Exception as e:
            raise Exception(f"خطأ في تصدير المحضر كملف DOCX: {str(e)}")

    @staticmethod
    def export_to_txt(meeting):
        """تصدير المحضر بصيغة نصية"""
        try:
            content = f"""عنوان الاجتماع: {meeting.title}
تاريخ الاجتماع: {meeting.created_at.strftime('%Y-%m-%d')}

محضر الاجتماع:
{meeting.minutes}

المهام:
"""
            # إضافة المهام
            for i, task in enumerate(meeting.tasks.all(), 1):
                assigned = f" (مسند إلى: {task.assigned_to})" if task.assigned_to else ""
                due_date = f" (تاريخ الاستحقاق: {task.due_date})" if task.due_date else ""
                content += f"{i}. {task.description}{assigned}{due_date}\n"

            content += "\nالقرارات:\n"

            # إضافة القرارات
            for i, decision in enumerate(meeting.decisions.all(), 1):
                decision_number = f" (رقم: {decision.decision_number})" if decision.decision_number else ""
                content += f"{i}. {decision.description}{decision_number}\n"

            return content.encode('utf-8')

        except Exception as e:
            raise Exception(f"خطأ في تصدير المحضر كملف نصي: {str(e)}")